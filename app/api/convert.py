from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import List, Optional, Tuple

from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import Response

from PIL import Image, ImageOps, ImageFilter
from docx import Document
from docx.shared import Pt

import pypdfium2 as pdfium

from app.core.config import settings
from app.core.public_limits import ensure_total_upload_size, ensure_upload_size

# Optional (best-effort) libraries. The service works without them, but PDF->DOCX quality
# is significantly better when they are installed.
try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover
    PdfReader = None

try:
    from pdf2docx import Converter as Pdf2DocxConverter  # type: ignore
except Exception:  # pragma: no cover
    Pdf2DocxConverter = None


router = APIRouter(prefix="/convert", tags=["convert"])

DEFAULT_RENDER_DPI = 220
OCR_RENDER_DPI = 300  # higher DPI improves OCR quality


# ----------------------------
# utilities
# ----------------------------

def _check_size(raw: bytes):
    ensure_upload_size(raw)


def _safe_base(filename: str) -> str:
    p = Path(filename)
    stem = p.stem or "file"
    out = "".join(c if (c.isalnum() or c in ("-", "_")) else "_" for c in stem)
    return (out[:80] or "file")


def _find_soffice() -> Optional[str]:
    return shutil.which("soffice") or shutil.which("libreoffice")


def _run_soffice_convert(input_path: Path, outdir: Path, target_ext: str, timeout_s: int = 180) -> Path:
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("LibreOffice not installed (soffice not found).")

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nodefault",
        "--norestore",
        "--convert-to", target_ext,
        "--outdir", str(outdir),
        str(input_path),
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout_s)
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out.")
    except subprocess.CalledProcessError as e:
        err = (e.stderr or b"").decode("utf-8", errors="ignore")[:2500]
        raise RuntimeError(f"LibreOffice conversion failed: {err}")

    expected = outdir / f"{input_path.stem}.{target_ext}"
    if expected.exists():
        return expected

    for f in outdir.glob(f"{input_path.stem}.*"):
        if f.suffix.lower() == f".{target_ext.lower()}":
            return f

    raise RuntimeError("Converted file not found.")


def _zip_bytes(files: List[Tuple[str, bytes]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for name, data in files:
            z.writestr(name, data)
    return buf.getvalue()


def _append_docx(dst: Document, src: Document) -> None:
    """
    Append body elements from src into dst.
    Keeps content editable and preserves most formatting.
    """
    for element in src.element.body:
        dst.element.body.append(deepcopy(element))


# ----------------------------
# PDF text detection (robust)
# ----------------------------

def _extract_text_pypdf(pdf_path: Path, page_index: int) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(pdf_path))
        if page_index < 0 or page_index >= len(reader.pages):
            return ""
        t = reader.pages[page_index].extract_text() or ""
        return " ".join(t.split())
    except Exception:
        return ""


def _extract_text_pdfium(pdf_path: Path, page_index: int, max_chars: int = 4000) -> str:
    """
    Fallback extraction via pdfium (often succeeds when pypdf is weak).
    """
    try:
        doc = pdfium.PdfDocument(str(pdf_path))
        page = doc[page_index]
        textpage = page.get_textpage()
        # get_text_range(start, count). If count=0 returns empty in some builds,
        # so we request a capped size.
        t = textpage.get_text_range(0, max_chars) or ""
        t = " ".join(str(t).split())
        textpage.close()
        page.close()
        doc.close()
        return t
    except Exception:
        return ""


def _page_has_text(pdf_path: Path, page_index: int, min_chars: int = 30) -> bool:
    """
    Decide if a page is "digital text" or "scan-like".
    """
    t = _extract_text_pypdf(pdf_path, page_index)
    if len(t) >= min_chars:
        return True
    t2 = _extract_text_pdfium(pdf_path, page_index)
    return len(t2) >= min_chars


def _pdf_has_extractable_text(pdf_path: Path, max_pages: int = 3, min_chars_total: int = 60) -> bool:
    """
    Quick heuristic for whole PDF (used as a fast path).
    """
    try:
        doc = pdfium.PdfDocument(str(pdf_path))
        n = min(len(doc), max_pages)
        doc.close()
    except Exception:
        return False

    total = 0
    for i in range(n):
        t = _extract_text_pypdf(pdf_path, i)
        if not t:
            t = _extract_text_pdfium(pdf_path, i)
        total += len(t)
        if total >= min_chars_total:
            return True
    return total >= min_chars_total


# ----------------------------
# PDF -> DOCX (best effort)
# ----------------------------

def _run_pdf2docx_convert(pdf_path: Path, out_docx_path: Path, start: Optional[int] = None, end: Optional[int] = None,
                          timeout_s: int = 240) -> None:
    """
    High-quality conversion for digital PDFs (layout + editability).
    start/end are 0-based, inclusive.
    """
    if Pdf2DocxConverter is None:
        raise RuntimeError("pdf2docx not installed")

    # pdf2docx runs in-process; to enforce a timeout we use a subprocess.
    # This avoids rare hangs on malformed PDFs.
    # pdf2docx Converter.convert(output, start=..., end=...)
    start_expr = "None" if start is None else str(int(start))
    end_expr = "None" if end is None else str(int(end))

    code = (
        "from pdf2docx import Converter; "
        f"c=Converter(r'{str(pdf_path)}'); "
        f"c.convert(r'{str(out_docx_path)}', start={start_expr}, end={end_expr}); "
        "c.close()"
    )
    cmd = ["python", "-c", code]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout_s)
    except subprocess.TimeoutExpired:
        raise RuntimeError("pdf2docx conversion timed out")
    except subprocess.CalledProcessError as e:
        err = (e.stderr or b"").decode("utf-8", errors="ignore")[:2500]
        raise RuntimeError(f"pdf2docx conversion failed: {err}")

    if not out_docx_path.exists() or out_docx_path.stat().st_size < 1024:
        raise RuntimeError("pdf2docx produced an empty output")


# ----------------------------
# pypdfium2 rendering
# ----------------------------

def _dpi_to_scale(dpi: int) -> float:
    # scale = dpi/72  because PDF uses points: 72pt = 1 inch
    dpi = max(72, min(int(dpi), 300))
    return dpi / 72.0


def _pdf_page_count(pdf_path: Path) -> int:
    doc = pdfium.PdfDocument(str(pdf_path))
    n = len(doc)
    doc.close()
    return n


def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """
    Light, safe preprocessing to improve OCR on faint scans.
    Keeps it fully open-source and fast.
    """
    # convert to RGB first
    if img.mode != "RGB":
        img = img.convert("RGB")

    # autocontrast helps washed-out scans
    img = ImageOps.autocontrast(img)

    # slight sharpening helps text edges
    img = img.filter(ImageFilter.SHARPEN)

    return img


def _render_pdf_page_pil(pdf_path: Path, page_index: int, dpi: int, preprocess_ocr: bool = False) -> Image.Image:
    scale = _dpi_to_scale(dpi)
    doc = pdfium.PdfDocument(str(pdf_path))
    page = doc[page_index]
    bitmap = page.render(scale=scale)
    img = bitmap.to_pil()  # RGBA or RGB depending on build
    # force white background
    if img.mode == "RGBA":
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    else:
        img = img.convert("RGB")

    if preprocess_ocr:
        img = _preprocess_for_ocr(img)

    page.close()
    doc.close()
    return img


def _render_pdf_to_images(pdf_path: Path, out_dir: Path, dpi: int, preprocess_ocr: bool = False) -> List[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    n = _pdf_page_count(pdf_path)
    paths: List[Path] = []
    for i in range(n):
        img = _render_pdf_page_pil(pdf_path, i, dpi=dpi, preprocess_ocr=preprocess_ocr)
        p = out_dir / f"page_{i+1:03d}.png"
        img.save(p, format="PNG", optimize=True)
        paths.append(p)
    return paths


def _render_pdf_pages_subset_to_images(pdf_path: Path, pages_1based: List[int], out_dir: Path,
                                       dpi: int, preprocess_ocr: bool = False) -> List[Path]:
    """
    Render a subset of pages (1-based page numbers) to images.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    paths: List[Path] = []
    for p in pages_1based:
        idx = p - 1
        img = _render_pdf_page_pil(pdf_path, idx, dpi=dpi, preprocess_ocr=preprocess_ocr)
        out_path = out_dir / f"page_{p:03d}.png"
        img.save(out_path, format="PNG", optimize=True)
        paths.append(out_path)
    return paths


def _render_pdf_all_pages_bytes(pdf_path: Path, dpi: int, fmt: str) -> List[Tuple[int, bytes]]:
    n = _pdf_page_count(pdf_path)
    out: List[Tuple[int, bytes]] = []
    for i in range(n):
        img = _render_pdf_page_pil(pdf_path, i, dpi=dpi)
        buf = io.BytesIO()
        if fmt == "png":
            img.save(buf, format="PNG", optimize=True)
        elif fmt == "jpg":
            img.save(buf, format="JPEG", quality=95, optimize=True)
        else:
            raise ValueError("fmt must be png or jpg")
        out.append((i + 1, buf.getvalue()))
    return out


# ----------------------------
# PaddleOCR v3 CLI (pp_structurev3)
# ----------------------------

def _run_paddle_pp_structurev3(img_dir: Path, out_dir: Path, lang: str = "en") -> Path:
    """
    Uses:
      paddleocr pp_structurev3 ...

    OCR recovery from images => generates docx in output folder.
    We try a couple flag styles because PaddleOCR CLI can differ by version.
    """
    paddle = shutil.which("paddleocr")
    if not paddle:
        raise RuntimeError("paddleocr not found in PATH (install in same venv).")

    out_dir.mkdir(parents=True, exist_ok=True)

    # Try variants for recovery flag (some builds accept true/True)
    cmd_variants = [
        [
            paddle,
            "pp_structurev3",
            f"--image_dir={str(img_dir)}",
            f"--output={str(out_dir)}",
            f"--lang={lang}",
            "--recovery=true",
        ],
        [
            paddle,
            "pp_structurev3",
            f"--image_dir={str(img_dir)}",
            f"--output={str(out_dir)}",
            f"--lang={lang}",
            "--recovery=True",
        ],
    ]

    last_err = ""
    for cmd in cmd_variants:
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=900)
            # Find latest docx
            candidates = sorted(out_dir.rglob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
            if candidates:
                return candidates[0]
            last_err = "PaddleOCR ran but no .docx output found."
        except subprocess.TimeoutExpired:
            last_err = "PaddleOCR conversion timed out."
        except subprocess.CalledProcessError as e:
            last_err = (e.stderr or b"").decode("utf-8", errors="ignore")[:3500]

    raise RuntimeError(f"PaddleOCR failed: {last_err}")


def _visual_docx_from_images(image_paths: List[Path]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(18)
    section.bottom_margin = Pt(18)
    section.left_margin = Pt(18)
    section.right_margin = Pt(18)

    usable_width = section.page_width - section.left_margin - section.right_margin

    for idx, img_path in enumerate(image_paths):
        doc.add_picture(str(img_path), width=usable_width)
        if idx != len(image_paths) - 1:
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ----------------------------
# image helpers
# ----------------------------

def _validate_image(file: UploadFile, kind: str):
    ext = (Path(file.filename or "").suffix or "").lower()
    ct = (file.content_type or "").lower()

    if kind == "png":
        if ct != "image/png" and ext != ".png":
            raise HTTPException(status_code=400, detail="Upload a PNG file")
    elif kind == "jpg":
        if ct not in ("image/jpeg", "image/jpg") and ext not in (".jpg", ".jpeg"):
            raise HTTPException(status_code=400, detail="Upload a JPG/JPEG file")
    else:
        raise ValueError("kind must be png or jpg")


async def _images_to_pdf(files: List[UploadFile]) -> bytes:
    if not files:
        raise HTTPException(status_code=400, detail="No images uploaded")

    total = 0
    imgs: List[Image.Image] = []
    for f in files:
        raw = await f.read()
        total += len(raw)
        ensure_total_upload_size(total)
        try:
            img = Image.open(io.BytesIO(raw))
            img = ImageOps.exif_transpose(img)
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                rgba = img.convert("RGBA")
                bg.paste(rgba, mask=rgba.split()[-1])
                img = bg
            else:
                img = img.convert("RGB")
            imgs.append(img)
        except Exception:
            raise HTTPException(status_code=400, detail=f"Invalid image: {f.filename}")

    buf = io.BytesIO()
    first, rest = imgs[0], imgs[1:]
    first.save(buf, format="PDF", save_all=True, append_images=rest)
    return buf.getvalue()


async def _image_to_word(file: UploadFile) -> bytes:
    raw = await file.read()
    _check_size(raw)
    try:
        img = Image.open(io.BytesIO(raw))
        img = ImageOps.exif_transpose(img).convert("RGB")
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG", optimize=True)
        img_buf.seek(0)
    except Exception:
        raise HTTPException(status_code=400, detail=f"Invalid image: {file.filename}")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(18)
    section.bottom_margin = Pt(18)
    section.left_margin = Pt(18)
    section.right_margin = Pt(18)
    usable_width = section.page_width - section.left_margin - section.right_margin

    doc.add_picture(img_buf, width=usable_width)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ----------------------------
# PDF -> DOCX improved pipeline (mixed PDFs supported)
# ----------------------------

def _contiguous_ranges(indices_0based: List[int]) -> List[Tuple[int, int]]:
    """
    Convert sorted indices into contiguous ranges (start,end) inclusive.
    """
    if not indices_0based:
        return []
    indices_0based = sorted(indices_0based)
    ranges: List[Tuple[int, int]] = []
    s = e = indices_0based[0]
    for x in indices_0based[1:]:
        if x == e + 1:
            e = x
        else:
            ranges.append((s, e))
            s = e = x
    ranges.append((s, e))
    return ranges


def _convert_pdf_to_docx_public_fast(pdf_path: Path, base_out_name: str) -> bytes:
    """
    Bounded conversion for the public free tools phase.
    It first attempts editable conversion, then quickly falls back to a visual DOCX
    so users do not sit at 90% for several minutes on Render Free.
    """
    n_pages = _pdf_page_count(pdf_path)
    max_pages = max(1, settings.PUBLIC_PDF_TO_WORD_MAX_PAGES)
    if n_pages > max_pages:
        raise HTTPException(
            status_code=413,
            detail=f"PDF has {n_pages} pages. Free public PDF to Word currently supports up to {max_pages} pages.",
        )

    timeout_s = max(10, settings.PUBLIC_PDF_TO_WORD_TIMEOUT_SECONDS)
    if _pdf_has_extractable_text(pdf_path) and Pdf2DocxConverter is not None:
        try:
            out_docx = pdf_path.parent / f"{base_out_name}_pdf2docx.docx"
            _run_pdf2docx_convert(pdf_path, out_docx, start=None, end=None, timeout_s=timeout_s)
            return out_docx.read_bytes()
        except Exception:
            pass

    img_dir = pdf_path.parent / "pages_visual_fast"
    dpi = max(96, settings.PUBLIC_PDF_RENDER_DPI)
    page_imgs = _render_pdf_to_images(pdf_path, img_dir, dpi=dpi, preprocess_ocr=False)
    return _visual_docx_from_images(page_imgs)


def _convert_pdf_to_docx_best(pdf_path: Path, base_out_name: str) -> bytes:
    """
    Best-effort conversion:
      1) If fully digital -> pdf2docx whole file
      2) Else per-page classify:
         - digital pages -> pdf2docx page ranges
         - scanned pages -> PaddleOCR on rendered images (high DPI + preprocess)
         - merge segments in order
      3) If anything fails badly -> visual DOCX fallback
    """
    n_pages = _pdf_page_count(pdf_path)

    # Fast path: if looks digital overall, try pdf2docx whole document first.
    if _pdf_has_extractable_text(pdf_path) and Pdf2DocxConverter is not None:
        try:
            out_docx = pdf_path.parent / f"{base_out_name}_pdf2docx.docx"
            _run_pdf2docx_convert(pdf_path, out_docx, start=None, end=None, timeout_s=300)
            return out_docx.read_bytes()
        except Exception:
            pass

    # Per-page classification (mixed PDFs)
    text_pages: List[int] = []
    scan_pages: List[int] = []
    for i in range(n_pages):
        if _page_has_text(pdf_path, i):
            text_pages.append(i)
        else:
            scan_pages.append(i)

    # If everything is scanned OR pdf2docx missing, go OCR batch
    if not text_pages or Pdf2DocxConverter is None:
        img_dir = pdf_path.parent / "pages_all"
        page_imgs = _render_pdf_to_images(pdf_path, img_dir, dpi=OCR_RENDER_DPI, preprocess_ocr=True)
        try:
            out_dir = pdf_path.parent / "paddle_out_all"
            docx_path = _run_paddle_pp_structurev3(img_dir, out_dir, lang="en")
            return docx_path.read_bytes()
        except Exception:
            return _visual_docx_from_images(page_imgs)

    # Otherwise: mixed or mostly digital -> build final docx by merging segments in order
    final_doc = Document()
    # Remove default empty paragraph if present
    if final_doc.paragraphs and not final_doc.paragraphs[0].text:
        p = final_doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Prepare segments in order: ranges of text pages and ranges of scan pages
    # We'll create a list of (kind, start0, end0)
    text_ranges = _contiguous_ranges(text_pages)
    scan_ranges = _contiguous_ranges(scan_pages)

    segments: List[Tuple[str, int, int]] = []
    ti = si = 0
    while ti < len(text_ranges) or si < len(scan_ranges):
        next_text = text_ranges[ti] if ti < len(text_ranges) else None
        next_scan = scan_ranges[si] if si < len(scan_ranges) else None

        if next_scan is None:
            segments.append(("text", next_text[0], next_text[1]))
            ti += 1
        elif next_text is None:
            segments.append(("scan", next_scan[0], next_scan[1]))
            si += 1
        else:
            # pick whichever starts earlier
            if next_text[0] <= next_scan[0]:
                segments.append(("text", next_text[0], next_text[1]))
                ti += 1
            else:
                segments.append(("scan", next_scan[0], next_scan[1]))
                si += 1

    try:
        for seg_idx, (kind, start0, end0) in enumerate(segments):
            if kind == "text":
                tmp_docx = pdf_path.parent / f"seg_text_{start0+1:03d}_{end0+1:03d}.docx"
                _run_pdf2docx_convert(pdf_path, tmp_docx, start=start0, end=end0, timeout_s=300)
                seg_doc = Document(str(tmp_docx))
                _append_docx(final_doc, seg_doc)
            else:
                # OCR only these pages
                pages_1based = list(range(start0 + 1, end0 + 2))
                img_dir = pdf_path.parent / f"seg_scan_imgs_{start0+1:03d}_{end0+1:03d}"
                _render_pdf_pages_subset_to_images(
                    pdf_path,
                    pages_1based=pages_1based,
                    out_dir=img_dir,
                    dpi=OCR_RENDER_DPI,
                    preprocess_ocr=True,
                )
                out_dir = pdf_path.parent / f"seg_scan_out_{start0+1:03d}_{end0+1:03d}"
                try:
                    docx_path = _run_paddle_pp_structurev3(img_dir, out_dir, lang="en")
                    seg_doc = Document(str(docx_path))
                    _append_docx(final_doc, seg_doc)
                except Exception:
                    # If OCR fails for this segment, insert as images (preserve geometry)
                    imgs = sorted(img_dir.glob("*.png"))
                    # Add page break before inserting images if needed
                    if final_doc.paragraphs:
                        final_doc.add_page_break()
                    # Insert images
                    section = final_doc.sections[-1]
                    usable_width = section.page_width - section.left_margin - section.right_margin
                    for j, imgp in enumerate(imgs):
                        final_doc.add_picture(str(imgp), width=usable_width)
                        if j != len(imgs) - 1:
                            final_doc.add_page_break()

            # Put a page break between segments (except after last)
            if seg_idx != len(segments) - 1:
                final_doc.add_page_break()

        out_buf = io.BytesIO()
        final_doc.save(out_buf)
        data = out_buf.getvalue()
        if len(data) > 1024:
            return data
    except Exception:
        pass

    # Last resort: visual DOCX of entire PDF
    img_dir = pdf_path.parent / "pages_fallback"
    page_imgs = _render_pdf_to_images(pdf_path, img_dir, dpi=DEFAULT_RENDER_DPI, preprocess_ocr=False)
    return _visual_docx_from_images(page_imgs)


# ----------------------------
# endpoints
# ----------------------------

@router.post("/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    pdf_bytes = await file.read()
    _check_size(pdf_bytes)

    if (Path(file.filename or "").suffix or "").lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Upload a .pdf file")

    base = _safe_base(file.filename or "document.pdf")
    out_name = f"{base}_converted.docx"

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        pdf_path = td / "input.pdf"
        pdf_path.write_bytes(pdf_bytes)

        if settings.PUBLIC_SITE_PHASE == "tools":
            docx_bytes = _convert_pdf_to_docx_public_fast(pdf_path, base_out_name=base)
        else:
            docx_bytes = _convert_pdf_to_docx_best(pdf_path, base_out_name=base)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


@router.post("/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)

    if (Path(file.filename or "").suffix or "").lower() != ".docx":
        raise HTTPException(status_code=400, detail="Upload a .docx file")

    base = _safe_base(file.filename or "document.docx")
    out_name = f"{base}_converted.pdf"

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        in_path = td / "input.docx"
        in_path.write_bytes(raw)
        pdf_path = _run_soffice_convert(in_path, td, "pdf", timeout_s=180)
        out_bytes = pdf_path.read_bytes()

    return Response(
        content=out_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


@router.post("/pdf-to-png")
async def pdf_to_png(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)
    if (Path(file.filename or "").suffix or "").lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Upload a .pdf file")

    base = _safe_base(file.filename or "document.pdf")

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        pdf_path = td / "input.pdf"
        pdf_path.write_bytes(raw)
        images = _render_pdf_all_pages_bytes(pdf_path, dpi=DEFAULT_RENDER_DPI, fmt="png")

    if len(images) == 1:
        p, data = images[0]
        return Response(
            content=data,
            media_type="image/png",
            headers={"Content-Disposition": f'attachment; filename="{base}_page_{p}.png"'},
        )

    zip_data = _zip_bytes([(f"{base}_page_{p}.png", d) for p, d in images])
    return Response(
        content=zip_data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}_png_images.zip"'},
    )


@router.post("/pdf-to-jpg")
async def pdf_to_jpg(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)
    if (Path(file.filename or "").suffix or "").lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Upload a .pdf file")

    base = _safe_base(file.filename or "document.pdf")

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        pdf_path = td / "input.pdf"
        pdf_path.write_bytes(raw)
        images = _render_pdf_all_pages_bytes(pdf_path, dpi=DEFAULT_RENDER_DPI, fmt="jpg")

    if len(images) == 1:
        p, data = images[0]
        return Response(
            content=data,
            media_type="image/jpeg",
            headers={"Content-Disposition": f'attachment; filename="{base}_page_{p}.jpg"'},
        )

    zip_data = _zip_bytes([(f"{base}_page_{p}.jpg", d) for p, d in images])
    return Response(
        content=zip_data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}_jpg_images.zip"'},
    )


@router.post("/word-to-png")
async def word_to_png(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)
    if (Path(file.filename or "").suffix or "").lower() != ".docx":
        raise HTTPException(status_code=400, detail="Upload a .docx file")

    base = _safe_base(file.filename or "document.docx")

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        docx_path = td / "input.docx"
        docx_path.write_bytes(raw)
        pdf_path = _run_soffice_convert(docx_path, td, "pdf", timeout_s=180)
        images = _render_pdf_all_pages_bytes(pdf_path, dpi=DEFAULT_RENDER_DPI, fmt="png")

    if len(images) == 1:
        p, data = images[0]
        return Response(
            content=data,
            media_type="image/png",
            headers={"Content-Disposition": f'attachment; filename="{base}_page_{p}.png"'},
        )

    zip_data = _zip_bytes([(f"{base}_page_{p}.png", d) for p, d in images])
    return Response(
        content=zip_data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}_png_images.zip"'},
    )


@router.post("/word-to-jpg")
async def word_to_jpg(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)
    if (Path(file.filename or "").suffix or "").lower() != ".docx":
        raise HTTPException(status_code=400, detail="Upload a .docx file")

    base = _safe_base(file.filename or "document.docx")

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        docx_path = td / "input.docx"
        docx_path.write_bytes(raw)
        pdf_path = _run_soffice_convert(docx_path, td, "pdf", timeout_s=180)
        images = _render_pdf_all_pages_bytes(pdf_path, dpi=DEFAULT_RENDER_DPI, fmt="jpg")

    if len(images) == 1:
        p, data = images[0]
        return Response(
            content=data,
            media_type="image/jpeg",
            headers={"Content-Disposition": f'attachment; filename="{base}_page_{p}.jpg"'},
        )

    zip_data = _zip_bytes([(f"{base}_page_{p}.jpg", d) for p, d in images])
    return Response(
        content=zip_data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}_jpg_images.zip"'},
    )


@router.post("/png-to-pdf")
async def png_to_pdf(files: List[UploadFile] = File(...)):
    for f in files:
        _validate_image(f, "png")
    pdf_bytes = await _images_to_pdf(files)
    base = _safe_base(files[0].filename or "image.png")
    out_name = f"{base}_converted.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


@router.post("/jpg-to-pdf")
async def jpg_to_pdf(files: List[UploadFile] = File(...)):
    for f in files:
        _validate_image(f, "jpg")
    pdf_bytes = await _images_to_pdf(files)
    base = _safe_base(files[0].filename or "image.jpg")
    out_name = f"{base}_converted.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


@router.post("/png-to-word")
async def png_to_word(file: UploadFile = File(...)):
    _validate_image(file, "png")
    out_bytes = await _image_to_word(file)
    base = _safe_base(file.filename or "image.png")
    out_name = f"{base}_converted.docx"
    return Response(
        content=out_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


@router.post("/jpg-to-word")
async def jpg_to_word(file: UploadFile = File(...)):
    _validate_image(file, "jpg")
    out_bytes = await _image_to_word(file)
    base = _safe_base(file.filename or "image.jpg")
    out_name = f"{base}_converted.docx"
    return Response(
        content=out_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )
